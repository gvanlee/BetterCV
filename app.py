from flask import Flask, render_template, request, redirect, url_for, flash, session, Response
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from database import get_db_connection, init_database, get_skill_categories
from datetime import datetime, timedelta
from authlib.integrations.flask_client import OAuth
from dotenv import load_dotenv
import json
import os
import uuid
import sqlite3
import hashlib
import secrets
import smtplib
from email.message import EmailMessage
from google import genai
from groq import Groq
import PyPDF2
import docx
import io
import tempfile
from ai_prompts import get_cv_parse_prompt, get_assignment_match_prompt

# Import authentication modules
from auth import User, get_all_users, get_all_whitelist, add_to_whitelist, remove_from_whitelist, update_user_role, deactivate_user, get_all_consultants
from decorators import admin_required, consultant_or_admin_required, check_consultant_access

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY', 'h9sdf696f34rhfvvhkxjvodfyg8yer89g7ye-8yhoiuver8v8erf98yyh34f')
app.permanent_session_lifetime = timedelta(hours=24)

# Initialize Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Please log in to access this page.'
login_manager.login_message_category = 'error'

# Initialize OAuth
oauth = OAuth(app)


def env_flag(name, default=False):
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() in {'1', 'true', 'yes', 'on'}


GOOGLE_OAUTH_ENABLED = bool(os.getenv('GOOGLE_CLIENT_ID', '').strip())
MICROSOFT_OAUTH_ENABLED = bool(os.getenv('MICROSOFT_CLIENT_ID', '').strip())
SMTP_HOST = os.getenv('SMTP_HOST', '').strip()
SMTP_PORT = int(os.getenv('SMTP_PORT', '587'))
SMTP_USERNAME = os.getenv('SMTP_USERNAME', '').strip()
SMTP_PASSWORD = os.getenv('SMTP_PASSWORD', '')
SMTP_USE_TLS = env_flag('SMTP_USE_TLS', True)
SMTP_USE_SSL = env_flag('SMTP_USE_SSL', False)
SMTP_FROM_EMAIL = os.getenv('SMTP_FROM_EMAIL', '').strip()
SMTP_FROM_NAME = os.getenv('SMTP_FROM_NAME', 'BetterCV').strip()
EMAIL_LOGIN_TOKEN_TTL_MINUTES = int(os.getenv('EMAIL_LOGIN_TOKEN_TTL_MINUTES', '30'))


def email_login_enabled():
    return bool(SMTP_HOST and SMTP_FROM_EMAIL)


def complete_login(user):
    login_user(user)
    session.permanent = True

    if user.consultant_id:
        session['consultant_id'] = user.consultant_id
    else:
        session.pop('consultant_id', None)


def create_email_login_token(user_id):
    conn = get_db_connection()
    now = datetime.utcnow()
    expires_at = now + timedelta(minutes=EMAIL_LOGIN_TOKEN_TTL_MINUTES)

    conn.execute(
        'DELETE FROM email_login_tokens WHERE used_at IS NOT NULL OR expires_at < ?',
        (now.isoformat(),)
    )

    raw_token = secrets.token_urlsafe(32)
    token_hash = hashlib.sha256(raw_token.encode('utf-8')).hexdigest()

    conn.execute(
        '''
            INSERT INTO email_login_tokens (user_id, token_hash, expires_at)
            VALUES (?, ?, ?)
        ''',
        (user_id, token_hash, expires_at.isoformat())
    )
    conn.commit()
    conn.close()

    return raw_token


def consume_email_login_token(raw_token):
    if not raw_token:
        return None

    token_hash = hashlib.sha256(raw_token.encode('utf-8')).hexdigest()
    conn = get_db_connection()
    token_row = conn.execute(
        '''
            SELECT id, user_id, expires_at, used_at
            FROM email_login_tokens
            WHERE token_hash = ?
        ''',
        (token_hash,)
    ).fetchone()

    if not token_row or token_row['used_at']:
        conn.close()
        return None

    expires_at = datetime.fromisoformat(token_row['expires_at'])
    if expires_at < datetime.utcnow():
        conn.close()
        return None

    cursor = conn.execute(
        'UPDATE email_login_tokens SET used_at = ? WHERE id = ? AND used_at IS NULL',
        (datetime.utcnow().isoformat(), token_row['id'])
    )
    conn.commit()
    conn.close()

    if cursor.rowcount != 1:
        return None

    return token_row['user_id']


def send_magic_link_email(recipient_email, login_link):
    import logging
    logger = logging.getLogger(__name__)
    logging.basicConfig(filename='bettercv.log', encoding='utf-8', level=logging.DEBUG)

    logger.debug(f"[EMAIL] send_magic_link_email called for: {recipient_email}")
    logger.debug(f"[EMAIL] email_login_enabled: {email_login_enabled()}")
    logger.debug(f"[EMAIL] SMTP_HOST: {SMTP_HOST}")
    logger.debug(f"[EMAIL] SMTP_PORT: {SMTP_PORT}")
    logger.debug(f"[EMAIL] SMTP_FROM_EMAIL: {SMTP_FROM_EMAIL}")
    logger.debug(f"[EMAIL] SMTP_USE_TLS: {SMTP_USE_TLS}")
    logger.debug(f"[EMAIL] SMTP_USE_SSL: {SMTP_USE_SSL}")
    logger.debug(f"[EMAIL] SMTP_USERNAME set: {bool(SMTP_USERNAME)}")
    logger.debug(f"[EMAIL] SMTP_PASSWORD set: {bool(SMTP_PASSWORD)}")

    if not email_login_enabled():
        logger.error("[EMAIL] Email login not enabled - SMTP_HOST or SMTP_FROM_EMAIL missing")
        return False

    message = EmailMessage()
    message['Subject'] = 'Your BetterCV sign-in link'
    message['From'] = f'{SMTP_FROM_NAME} <{SMTP_FROM_EMAIL}>' if SMTP_FROM_NAME else SMTP_FROM_EMAIL
    message['To'] = recipient_email
    message.set_content(
        f"Use this one-time link to sign in to BetterCV:\n\n{login_link}\n\n"
        f"This link expires in {EMAIL_LOGIN_TOKEN_TTL_MINUTES} minutes and can only be used once."
    )
    logger.debug(f"[EMAIL] Email message created - From: {message['From']}, To: {message['To']}")

    try:
        logger.debug(f"[EMAIL] Attempting SMTP connection: host={SMTP_HOST}, port={SMTP_PORT}, ssl={SMTP_USE_SSL}, tls={SMTP_USE_TLS}")

        if SMTP_USE_SSL:
            logger.debug("[EMAIL] Using SMTP_SSL connection")
            with smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT, timeout=20) as smtp:
                logger.debug("[EMAIL] SMTP_SSL connection established")
                if SMTP_USERNAME:
                    logger.debug(f"[EMAIL] Authenticating as: {SMTP_USERNAME}")
                    smtp.login(SMTP_USERNAME, SMTP_PASSWORD)
                    logger.debug("[EMAIL] Authentication successful")
                else:
                    logger.debug("[EMAIL] No authentication required (no username)")
                logger.debug("[EMAIL] Sending message...")
                smtp.send_message(message)
                logger.info(f"[EMAIL] Message sent successfully to {recipient_email}")
        else:
            logger.debug("[EMAIL] Using standard SMTP connection")
            with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=20) as smtp:
                logger.debug("[EMAIL] SMTP connection established")
                if SMTP_USE_TLS:
                    logger.debug("[EMAIL] Starting TLS...")
                    smtp.starttls()
                    logger.debug("[EMAIL] TLS started")
                else:
                    logger.debug("[EMAIL] TLS not enabled")

                if SMTP_USERNAME:
                    logger.debug(f"[EMAIL] Authenticating as: {SMTP_USERNAME}")
                    smtp.login(SMTP_USERNAME, SMTP_PASSWORD)
                    logger.debug("[EMAIL] Authentication successful")
                else:
                    logger.debug("[EMAIL] No authentication required (no username)")

                logger.debug("[EMAIL] Sending message...")
                smtp.send_message(message)
                logger.info(f"[EMAIL] Message sent successfully to {recipient_email}")

        return True
    except smtplib.SMTPAuthenticationError as auth_err:
        logger.error(f"[EMAIL] SMTP Authentication failed: {str(auth_err)}")
        return False
    except smtplib.SMTPException as smtp_err:
        logger.error(f"[EMAIL] SMTP error: {str(smtp_err)}")
        return False
    except OSError as os_err:
        logger.error(f"[EMAIL] Connection error (network/DNS): {str(os_err)}")
        return False
    except Exception as e:
        logger.error(f"[EMAIL] Unexpected error sending magic-link email: {str(e)}")
        import traceback
        logger.error(f"[EMAIL] Traceback: {traceback.format_exc()}")
        return False

# Configure Google OAuth
google = oauth.register(
    name='google',
    client_id=os.getenv('GOOGLE_CLIENT_ID'),
    client_secret=os.getenv('GOOGLE_CLIENT_SECRET'),
    server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
    client_kwargs={'scope': 'openid email profile'}
)

# Configure Microsoft OAuth
microsoft = oauth.register(
    name='microsoft',
    client_id=os.getenv('MICROSOFT_CLIENT_ID'),
    client_secret=os.getenv('MICROSOFT_CLIENT_SECRET'),
    # authorize_url='https://login.microsoftonline.com/organizations/oauth2/v2.0/authorize',
    authorize_url='https://login.microsoftonline.com/39e6a0f2-2da0-42b2-8a33-476485f72038/oauth2/v2.0/authorize',
    authorize_params=None,
    # access_token_url='https://login.microsoftonline.com/organizations/oauth2/v2.0/token',
    access_token_url='https://login.microsoftonline.com/39e6a0f2-2da0-42b2-8a33-476485f72038/oauth2/v2.0/token',
    access_token_params=None,
    refresh_token_url=None,
    client_kwargs={
        'scope': 'openid email profile',
         'validate_iss': False
    },
    jwks_uri="https://login.microsoftonline.com/common/discovery/v2.0/keys"    
)

@login_manager.user_loader
def load_user(user_id):
    """Load user by ID for Flask-Login"""
    return User.get_by_id(user_id)

# Load translations
TRANSLATIONS = {}
TRANSLATIONS_DIR = os.path.join(os.path.dirname(__file__), 'translations')

def load_translations():
    """Load all translation files."""
    for filename in os.listdir(TRANSLATIONS_DIR):
        if filename.endswith('.json'):
            lang_code = filename[:-5]  # Remove .json extension
            with open(os.path.join(TRANSLATIONS_DIR, filename), 'r', encoding='utf-8') as f:
                TRANSLATIONS[lang_code] = json.load(f)

load_translations()

def get_translation(key, lang=None):
    """Get translation for a key in the current language."""
    if lang is None:
        lang = session.get('language', 'en')
    
    keys = key.split('.')
    value = TRANSLATIONS.get(lang, TRANSLATIONS['en'])
    
    for k in keys:
        if isinstance(value, dict):
            value = value.get(k, key)
        else:
            return key
    
    return value

def format_markdown_text(text):
    """Convert simple markdown (lines starting with *) to HTML."""
    if not text:
        return text
    
    lines = text.strip().split('\n')
    result = []
    in_list = False
    
    for line in lines:
        line = line.strip()
        if not line:
            if in_list:
                result.append('</ul>')
                in_list = False
            result.append('<br>')
            continue
        
        if line.startswith('* '):
            if not in_list:
                result.append('<ul>')
                in_list = True
            result.append(f'<li>{line[2:]}</li>')
        else:
            if in_list:
                result.append('</ul>')
                in_list = False
            result.append(f'<p>{line}</p>')
    
    if in_list:
        result.append('</ul>')
    
    return '\n'.join(result)

def extract_list_items(text):
    """Extract bulleted list items from markdown text for Word export."""
    if not text:
        return [], []
    
    lines = text.strip().split('\n')
    list_items = []
    paragraph_items = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('* '):
            list_items.append(line[2:])
        else:
            paragraph_items.append(line)
    
    return list_items, paragraph_items

@app.context_processor
def inject_translations():
    """Make translation function and current language available to all templates."""
    lang = session.get('language', 'en')
    conn = get_db_connection()
    consultants = get_consultants(conn)
    current_consultant_id = resolve_current_consultant_id(conn)
    current_consultant = None
    if current_consultant_id is not None:
        current_consultant = conn.execute(
            'SELECT id, display_name FROM consultants WHERE id = ?',
            (current_consultant_id,)
        ).fetchone()
    conn.close()

    return {
        't': TRANSLATIONS.get(lang, TRANSLATIONS['en']),
        'current_lang': lang,
        'current_user': current_user,
        'consultants': consultants,
        'current_consultant': current_consultant,
        'available_languages': {
            'en': 'English',
            'nl': 'Nederlands'
        },
        'format_markdown_text': format_markdown_text
    }

# Register custom Jinja filters
app.jinja_env.filters['format_markdown'] = format_markdown_text

@app.route('/set-language/<lang>')
def set_language(lang):
    """Set the user's language preference."""
    if lang in TRANSLATIONS:
        session['language'] = lang
    return redirect(request.referrer or url_for('index'))


# ==================== Authentication Routes ====================

@app.route('/login')
def login():
    """Login page"""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    return render_template(
        'login.html',
        show_google_login=GOOGLE_OAUTH_ENABLED,
        show_microsoft_login=MICROSOFT_OAUTH_ENABLED,
        show_email_login=email_login_enabled()
    )


@app.route('/logout')
@login_required
def logout():
    """Logout user"""
    logout_user()
    flash(get_translation('messages.logged_out'), 'success')
    return redirect(url_for('login'))


@app.route('/auth/google')
def auth_google():
    """Initiate Google OAuth"""
    if not GOOGLE_OAUTH_ENABLED:
        flash(get_translation('messages.oauth_not_configured'), 'error')
        return redirect(url_for('login'))

    # redirect_uri = url_for('auth_google_callback', _external=True)
    redirect_uri = (os.getenv('BASE_URL') + os.getenv('GOOGLE_OAUTH_REDIRECT_URI')) or url_for('auth_google_callback', _external=True)

    return google.authorize_redirect(redirect_uri)


@app.route('/auth/google/callback')
def auth_google_callback():
    """Google OAuth callback"""
    try:
        token = google.authorize_access_token()
        user_info = token.get('userinfo')
        
        if not user_info:
            flash(get_translation('messages.oauth_failed'), 'error')
            return redirect(url_for('login'))
        
        email = user_info.get('email')
        name = user_info.get('name')
        oauth_id = user_info.get('sub')
        
        # Check if user exists
        user = User.get_by_email(email)
        
        if not user:
            # Check whitelist and create user
            user = User.create_user(email, name, 'google', oauth_id)
            if not user:
                flash(get_translation('messages.not_whitelisted'), 'error')
                return redirect(url_for('login'))
        
        # Update last login
        User.update_last_login(user.id)
        
        # Log in user
        complete_login(user)
        flash(get_translation('messages.login_success'), 'success')
        
        return redirect(url_for('index'))
        
    except Exception as e:
        print(f"OAuth error: {str(e)}")
        flash(get_translation('messages.oauth_error'), 'error')
        return redirect(url_for('login'))


@app.route('/auth/microsoft')
def auth_microsoft():
    """Initiate Microsoft OAuth"""
    if not MICROSOFT_OAUTH_ENABLED:
        flash(get_translation('messages.oauth_not_configured'), 'error')
        return redirect(url_for('login'))

    # redirect_uri = url_for('auth_microsoft_callback', _external=True)
    redirect_uri = (os.getenv('BASE_URL') + os.getenv('MICROSOFT_OAUTH_REDIRECT_URI')) or url_for('auth_microsoft_callback', _external=True)
    print(f"Initiating Microsoft OAuth, redirect_uri: {redirect_uri}")
    return microsoft.authorize_redirect(redirect_uri)


@app.route('/auth/microsoft/callback')
def auth_microsoft_callback():
    """Microsoft OAuth callback"""
    try:
        token = microsoft.authorize_access_token()
        
        # Get user info from Microsoft Graph API
        resp = microsoft.get('https://graph.microsoft.com/v1.0/me', token=token)
        user_info = resp.json()
        
        email = user_info.get('mail') or user_info.get('userPrincipalName')
        name = user_info.get('displayName')
        oauth_id = user_info.get('id')
        
        if not email:
            print(f"Microsoft OAuth response missing email: {user_info}, json: {resp.text}")
            flash(get_translation('messages.oauth_failed'), 'error')
            return redirect(url_for('login'))
        
        # Check if user exists
        user = User.get_by_email(email)
        
        if not user:
            # Check whitelist and create user
            user = User.create_user(email, name, 'microsoft', oauth_id)
            if not user:
                flash(get_translation('messages.not_whitelisted'), 'error')
                return redirect(url_for('login'))
        
        # Update last login
        User.update_last_login(user.id)
        
        # Log in user
        complete_login(user)
        flash(get_translation('messages.login_success'), 'success')
        
        return redirect(url_for('index'))
        
    except Exception as e:
        print(f"OAuth error: {str(e)}")
        flash(get_translation('messages.oauth_error'), 'error')
        return redirect(url_for('login'))


@app.route('/auth/email/request', methods=['POST'])
def auth_email_request():
    """Send one-time login link to existing user by e-mail."""
    import logging
    logger = logging.getLogger(__name__)
    logging.basicConfig(filename='bettercv.log', encoding='utf-8', level=logging.DEBUG)

    logger.debug("[AUTH_EMAIL] auth_email_request called")

    if not email_login_enabled():
        logger.error("[AUTH_EMAIL] Email login not enabled")
        flash(get_translation('messages.email_login_not_configured'), 'error')
        return redirect(url_for('login'))

    email = (request.form.get('email') or '').strip().lower()
    logger.debug(f"[AUTH_EMAIL] Email requested: {email}")

    if not email:
        logger.warning("[AUTH_EMAIL] No email provided")
        flash(get_translation('messages.email_required'), 'error')
        return redirect(url_for('login'))

    user = User.get_by_email(email)
    logger.debug(f"[AUTH_EMAIL] User lookup result: {bool(user)}")

    if not user:
        logger.debug(f"[AUTH_EMAIL] No user found for {email}, checking whitelist...")
        whitelist_entry = User.is_whitelisted(email)
        
        if not whitelist_entry:
            domain = User._get_email_domain(email)
            auto_domains = User._get_auto_whitelist_domains()
            if domain and domain in auto_domains:
                logger.info(f"[AUTH_EMAIL] Domain {domain} is auto-whitelisted, adding to whitelist")
                from auth import add_to_whitelist
                add_to_whitelist(email, 'consultant', notes='Auto-whitelisted by domain (email login)')
                whitelist_entry = User.is_whitelisted(email)
            else:
                logger.debug(f"[AUTH_EMAIL] Email {email} and domain not whitelisted")
        
        if whitelist_entry:
            logger.info(f"[AUTH_EMAIL] Email {email} is whitelisted, creating user")
            user = User.create_user(email, None, 'email', None)
            if user:
                logger.info(f"[AUTH_EMAIL] User created for {email}, user_id={user.id}")
            else:
                logger.error(f"[AUTH_EMAIL] Failed to create user for {email}")
        else:
            logger.debug(f"[AUTH_EMAIL] Email {email} not whitelisted (security: not exposing this to user)")

    if user:
        logger.info(f"[AUTH_EMAIL] User ready for {email}, user_id={user.id}")
        token = create_email_login_token(user.id)
        logger.debug(f"[AUTH_EMAIL] Token created (hash only, not logged)")

        login_link = url_for('auth_email_verify', token=token, _external=True)
        logger.debug(f"[AUTH_EMAIL] Login link generated")

        sent = send_magic_link_email(email, login_link)
        logger.info(f"[AUTH_EMAIL] Email send result: {sent}")

        if not sent:
            logger.error(f"[AUTH_EMAIL] Failed to send email to {email}")
            flash(get_translation('messages.email_login_send_failed'), 'error')
            return redirect(url_for('login'))
        else:
            logger.info(f"[AUTH_EMAIL] Email successfully sent to {email}")
    else:
        logger.debug(f"[AUTH_EMAIL] User not eligible or not found for {email}")

    flash(get_translation('messages.email_login_link_sent'), 'success')
    logger.debug("[AUTH_EMAIL] Generic success message shown to user")
    return redirect(url_for('login'))


@app.route('/auth/email/verify')
def auth_email_verify():
    """Validate one-time login token and sign in user."""
    import logging
    logger = logging.getLogger(__name__)

    logger.debug("[AUTH_EMAIL_VERIFY] auth_email_verify called")

    token = (request.args.get('token') or '').strip()
    logger.debug(f"[AUTH_EMAIL_VERIFY] Token present: {bool(token)}")

    if not token:
        logger.warning("[AUTH_EMAIL_VERIFY] No token provided")
        flash(get_translation('messages.email_login_invalid_token'), 'error')
        return redirect(url_for('login'))

    user_id = consume_email_login_token(token)
    logger.debug(f"[AUTH_EMAIL_VERIFY] Token validation result: {bool(user_id)}")

    if not user_id:
        logger.warning("[AUTH_EMAIL_VERIFY] Invalid or expired token")
        flash(get_translation('messages.email_login_invalid_token'), 'error')
        return redirect(url_for('login'))

    user = User.get_by_id(user_id)
    logger.debug(f"[AUTH_EMAIL_VERIFY] User lookup result: {bool(user)}")

    if not user:
        logger.error(f"[AUTH_EMAIL_VERIFY] User not found for valid token user_id={user_id}")
        flash(get_translation('messages.email_login_invalid_token'), 'error')
        return redirect(url_for('login'))

    logger.info(f"[AUTH_EMAIL_VERIFY] Valid token for user: {user.email}")
    User.update_last_login(user.id)
    complete_login(user)
    flash(get_translation('messages.login_success'), 'success')
    logger.info(f"[AUTH_EMAIL_VERIFY] User {user.email} logged in successfully via email link")
    return redirect(url_for('index'))


# ==================== Admin Routes ====================

@app.route('/admin')
@admin_required
def admin_dashboard():
    """Admin dashboard"""
    users = get_all_users()
    whitelist = get_all_whitelist()
    skill_categories = get_skill_categories()
    return render_template('admin.html', 
                         users=users, 
                         whitelist=whitelist,
                         skill_categories=skill_categories)


@app.route('/consultants-management')
@admin_required
def consultants_management():
    """Consultants management page"""
    consultants = get_all_consultants()
    return render_template('consultants_management.html',
                         all_consultants=consultants)


@app.route('/admin/skill-categories/add', methods=['POST'])
@admin_required
def admin_add_skill_category():
    """Add new skill category."""
    name = request.form.get('name')
    
    if name:
        conn = get_db_connection()
        try:
            conn.execute(
                'INSERT INTO skill_categories (name) VALUES (?)',
                (name,)
            )
            conn.commit()
            flash(get_translation('messages.category_added'), 'success')
        except sqlite3.IntegrityError:
            flash(get_translation('messages.category_exists'), 'error')
        conn.close()
    
    return redirect(url_for('admin_dashboard'))


@app.route('/admin/skill-categories/delete/<int:category_id>', methods=['POST'])
@admin_required
def admin_delete_skill_category(category_id):
    """Delete skill category."""
    conn = get_db_connection()
    # Check if category is used in experience records
    experience_usage = conn.execute('''
        SELECT COUNT(*) 
        FROM experience_skills es 
        JOIN skills s ON es.skill_id = s.id 
        WHERE s.category_id = ?
    ''', (category_id,)).fetchone()[0]
    
    # Check if category is assigned to any skills (FK safety)
    skill_usage = conn.execute('SELECT COUNT(*) FROM skills WHERE category_id = ?', (category_id,)).fetchone()[0]
    
    if experience_usage > 0:
        flash(get_translation('messages.category_referenced_by_experience'), 'error')
    elif skill_usage > 0:
        flash(get_translation('messages.category_in_use'), 'error')
    else:
        conn.execute('DELETE FROM skill_categories WHERE id = ?', (category_id,))
        conn.commit()
        flash(get_translation('messages.category_deleted'), 'success')
    conn.close()
    return redirect(url_for('admin_dashboard'))


@app.route('/admin/whitelist/add', methods=['POST'])
@admin_required
def admin_add_whitelist():
    """Add email to whitelist"""
    email = request.form.get('email')
    role = request.form.get('role', 'consultant')
    notes = request.form.get('notes', '')
    
    success, message = add_to_whitelist(email, role, current_user.id, notes)
    
    if success:
        flash(get_translation('messages.whitelist_added'), 'success')
    else:
        flash(f"{get_translation('messages.whitelist_error')}: {message}", 'error')
    
    return redirect(url_for('admin_dashboard'))


@app.route('/admin/whitelist/remove/<int:whitelist_id>', methods=['POST'])
@admin_required
def admin_remove_whitelist(whitelist_id):
    """Remove email from whitelist"""
    remove_from_whitelist(whitelist_id)
    flash(get_translation('messages.whitelist_removed'), 'success')
    return redirect(url_for('admin_dashboard'))


@app.route('/admin/user/update-role/<int:user_id>', methods=['POST'])
@admin_required
def admin_update_user_role(user_id):
    """Update user role"""
    new_role = request.form.get('role')
    update_user_role(user_id, new_role)
    flash(get_translation('messages.user_role_updated'), 'success')
    return redirect(url_for('admin_dashboard'))


@app.route('/admin/user/deactivate/<int:user_id>', methods=['POST'])
@admin_required
def admin_deactivate_user(user_id):
    """Deactivate user account"""
    if user_id == current_user.id:
        flash(get_translation('messages.cannot_deactivate_self'), 'error')
    else:
        # Get the user being deactivated to check their role
        conn = get_db_connection()
        user = conn.execute('SELECT role FROM users WHERE id = ?', (user_id,)).fetchone()
        conn.close()
        
        if user and user['role'] == 'admin':
            flash(get_translation('messages.cannot_deactivate_admin'), 'error')
        else:
            deactivate_user(user_id)
            flash(get_translation('messages.user_deactivated'), 'success')
    return redirect(url_for('admin_dashboard'))


def get_consultants(conn):
    """Retrieve all consultants ordered by display name."""
    if current_user.is_authenticated and not current_user.is_admin():
        if not current_user.consultant_id:
            return []
        return conn.execute(
            'SELECT id, display_name FROM consultants WHERE id = ? ORDER BY display_name, id',
            (current_user.consultant_id,)
        ).fetchall()

    return conn.execute(
        'SELECT id, display_name FROM consultants ORDER BY display_name, id'
    ).fetchall()


def resolve_current_consultant_id(conn):
    if current_user.is_authenticated and not current_user.is_admin():
        if not current_user.consultant_id:
            return None

        own_consultant = conn.execute(
            'SELECT id FROM consultants WHERE id = ?',
            (current_user.consultant_id,)
        ).fetchone()
        if not own_consultant:
            return None

        session['consultant_id'] = current_user.consultant_id
        return current_user.consultant_id

    consultant_id = session.get('consultant_id')
    if consultant_id:
        existing = conn.execute(
            'SELECT id FROM consultants WHERE id = ?',
            (consultant_id,)
        ).fetchone()
        if existing:
            return consultant_id

    first = conn.execute(
        'SELECT id FROM consultants ORDER BY display_name, id LIMIT 1'
    ).fetchone()
    if first:
        session['consultant_id'] = first['id']
        return first['id']

    return None


@app.before_request
def ensure_consultant_selected():
    if request.endpoint in {
        'static',
        'set_language',
        'list_consultants',
        'add_consultant',
        'switch_consultant',
        'import_consultant',
        'export_consultant',
        'parse_cv',
        'assignment_match',
        'review_parsed_cv',
        'import_parsed_cv',
        'export_cv',
        'preview_cv',
        'export_cv_download',
        'consultants_management',
        'admin_dashboard',
        'admin_add_whitelist',
        'admin_remove_whitelist',
        'admin_update_user_role',
        'admin_deactivate_user',
        'admin_add_skill_category',
        'admin_delete_skill_category',
        'delete_consultant'
    }:
        return None

    if current_user.is_authenticated and not current_user.is_admin():
        if current_user.consultant_id:
            session['consultant_id'] = current_user.consultant_id
            return None
        flash(get_translation('messages.user_no_consultant_access'), 'error')
        logout_user()
        session.pop('consultant_id', None)
        return redirect(url_for('login'))

    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    conn.close()

    if consultant_id is None:
        return redirect(url_for('consultants_management'))

    return None


# ==================== AI Configuration and CV Parsing ====================

# Configure Gemini AI - Set your API key as environment variable
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
genai_client = None
if GEMINI_API_KEY:
    genai_client = genai.Client(api_key=GEMINI_API_KEY)

# Configure Groq AI
GROQ_API_KEY = os.getenv('GROQ_API_KEY')
groq_client = None
if GROQ_API_KEY:
    groq_client = Groq(api_key=GROQ_API_KEY)

# ==================== CV Parsing Data Storage ====================

def get_cv_temp_dir():
    """Get or create the temporary directory for CV parsing data."""
    temp_dir = os.path.join(tempfile.gettempdir(), 'bettercv_cv_parsing')
    os.makedirs(temp_dir, exist_ok=True)
    return temp_dir


def store_parsed_cv_data(parsed_data, cv_text_preview):
    """Store parsed CV data in a temporary file and return the ID."""
    cv_id = str(uuid.uuid4())
    temp_dir = get_cv_temp_dir()
    temp_file = os.path.join(temp_dir, f'{cv_id}.json')
    
    data_to_store = {
        'parsed_data': parsed_data,
        'cv_text_preview': cv_text_preview,
        'created_at': datetime.now().isoformat()
    }
    
    with open(temp_file, 'w', encoding='utf-8') as f:
        json.dump(data_to_store, f)
    
    return cv_id


def retrieve_parsed_cv_data(cv_id):
    """Retrieve parsed CV data from temporary file."""
    temp_dir = get_cv_temp_dir()
    temp_file = os.path.join(temp_dir, f'{cv_id}.json')
    
    if not os.path.exists(temp_file):
        return None, None
    
    try:
        with open(temp_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return data.get('parsed_data'), data.get('cv_text_preview')
    except (json.JSONDecodeError, IOError):
        return None, None


def delete_parsed_cv_data(cv_id):
    """Delete temporary CV parsing data file."""
    temp_dir = get_cv_temp_dir()
    temp_file = os.path.join(temp_dir, f'{cv_id}.json')
    
    try:
        if os.path.exists(temp_file):
            os.remove(temp_file)
    except OSError:
        pass  # Silently ignore if file can't be deleted


def extract_text_from_pdf(file_content):
    """Extract text from PDF file."""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        return f"Error reading PDF: {str(e)}"


def extract_text_from_docx(file_content):
    """Extract text from DOCX file."""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"


def extract_text_from_file(file):
    """Extract text from uploaded file based on file type."""
    filename = file.filename.lower()
    file_content = file.read()
    file.seek(0)  # Reset file pointer
    
    if filename.endswith('.pdf'):
        return extract_text_from_pdf(file_content)
    elif filename.endswith(('.docx', '.doc')):
        return extract_text_from_docx(file_content)
    elif filename.endswith('.txt'):
        return file_content.decode('utf-8', errors='ignore')
    else:
        return "Unsupported file type. Please upload PDF, DOCX, or TXT files."


def parse_cv_with_ai(cv_text, provider='gemini'):
    """Parse CV text using chosen AI provider and return structured data."""
    prompt = get_cv_parse_prompt(cv_text)
    
    if provider == 'gemini':
        if not GEMINI_API_KEY or genai_client is None:
            return None, "Gemini API key not configured."
        
        try:
            response = genai_client.models.generate_content(
                model='gemini-flash-latest',
                contents=prompt
            )
            if not response.text:
                return None, "No response from Gemini AI"
            json_text = response.text
        except Exception as e:
            return None, f"Gemini error: {str(e)}"
            
    elif provider == 'groq':
        if not GROQ_API_KEY or groq_client is None:
            return None, "Groq API key not configured."
            
        try:
            response = groq_client.chat.completions.create(
                model='llama-3.3-70b-versatile',
                messages=[
                    {"role": "system", "content": "You are a specialized CV parsing assistant. Always respond with valid JSON only."},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"}
            )
            json_text = response.choices[0].message.content
        except Exception as e:
            return None, f"Groq error: {str(e)}"
    else:
        return None, f"Unknown AI provider: {provider}"

    # Try to parse the JSON response
    try:
        # Clean up the response text (remove markdown formatting if present)
        json_text = json_text.strip()
        if json_text.startswith('```json'):
            json_text = json_text[7:]
        if json_text.endswith('```'):
            json_text = json_text[:-3]
        
        parsed_data = json.loads(json_text)
        return parsed_data, None
    except json.JSONDecodeError as e:
        return None, f"Failed to parse AI response as JSON: {str(e)}\nResponse: {json_text[:200]}..."


def match_assignment_with_ai(assignment_description, consultant_payloads, provider='gemini'):
    """Match assignment text against selected consultant profiles using AI."""
    consultants_json = json.dumps(consultant_payloads, indent=2, ensure_ascii=False)
    prompt = get_assignment_match_prompt(assignment_description, consultants_json)

    if provider == 'gemini':
        if not GEMINI_API_KEY or genai_client is None:
            return None, consultants_json, "Gemini API key not configured."

        try:
            response = genai_client.models.generate_content(
                model='gemini-flash-latest',
                contents=prompt
            )
            if not response.text:
                return None, consultants_json, "No response from Gemini AI"
            return response.text, consultants_json, None
        except Exception as e:
            return None, consultants_json, f"Gemini error: {str(e)}"

    if provider == 'groq':
        if not GROQ_API_KEY or groq_client is None:
            return None, consultants_json, "Groq API key not configured."

        try:
            response = groq_client.chat.completions.create(
                model='llama-3.3-70b-versatile',
                messages=[
                    {"role": "system", "content": "You are a recruitment matching assistant."},
                    {"role": "user", "content": prompt}
                ]
            )
            return response.choices[0].message.content, consultants_json, None
        except Exception as e:
            return None, consultants_json, f"Groq error: {str(e)}"

    return None, consultants_json, f"Unknown AI provider: {provider}"


def build_consultant_ai_payload(conn, consultant_id):
    """Build normalized consultant payload for export/matching AI input."""
    consultant = conn.execute(
        'SELECT id, display_name FROM consultants WHERE id = ?',
        (consultant_id,)
    ).fetchone()

    if not consultant:
        return None

    payload = {
        'consultant': {
            'id': consultant['id'],
            'display_name': consultant['display_name']
        },
        'personal_info': None,
        'work_experience': [],
        'education': [],
        'skills': [],
        'projects': [],
        'certifications': []
    }

    personal_info = conn.execute(
        'SELECT * FROM personal_info WHERE consultant_id = ? ORDER BY id DESC LIMIT 1',
        (consultant_id,)
    ).fetchone()
    if personal_info:
        payload['personal_info'] = dict(personal_info)
        payload['personal_info'].pop('id', None)
        payload['personal_info'].pop('consultant_id', None)
        payload['personal_info'].pop('created_at', None)
        payload['personal_info'].pop('updated_at', None)

    for table, key in [
        ('work_experience', 'work_experience'),
        ('education', 'education'),
        ('skills', 'skills'),
        ('projects', 'projects'),
        ('certifications', 'certifications')
    ]:
        rows = conn.execute(
            f'SELECT * FROM {table} WHERE consultant_id = ? ORDER BY id',
            (consultant_id,)
        ).fetchall()
        payload[key] = []
        for row in rows:
            row_data = dict(row)

            if table == 'skills' and row_data.get('category_id'):
                cat = conn.execute('SELECT name FROM skill_categories WHERE id = ?', (row_data['category_id'],)).fetchone()
                if cat:
                    row_data['category_name'] = cat['name']

            if table == 'work_experience':
                skills = conn.execute('''
                    SELECT s.skill_name
                    FROM skills s
                    JOIN experience_skills es ON s.id = es.skill_id
                    WHERE es.experience_id = ?
                ''', (row['id'],)).fetchall()
                row_data['skills'] = [s['skill_name'] for s in skills]

            if table == 'projects':
                skills = conn.execute('''
                    SELECT s.skill_name
                    FROM skills s
                    JOIN project_skills ps ON s.id = ps.skill_id
                    WHERE ps.project_id = ?
                ''', (row['id'],)).fetchall()
                row_data['skills'] = [s['skill_name'] for s in skills]

            if table == 'certifications' and row_data.get('issue_date'):
                issue_date_str = str(row_data['issue_date'])
                if issue_date_str and len(issue_date_str) >= 4:
                    row_data['issue_year'] = issue_date_str[:4]

            row_data.pop('id', None)
            row_data.pop('consultant_id', None)
            row_data.pop('created_at', None)
            row_data.pop('updated_at', None)
            row_data.pop('category', None)
            row_data.pop('years_of_experience', None)
            payload[key].append(row_data)

    return payload


def parse_cv_with_gemini(cv_text):
    """Backward compatibility for existing code."""
    return parse_cv_with_ai(cv_text, 'gemini')
            
@app.route('/')
@login_required
def index():
    """Home page with navigation to all sections."""
    return render_template('index.html')


# ==================== Consultant Routes ====================

@app.route('/consultants')
@login_required
def list_consultants():
    """Redirect to admin dashboard (consultants are managed there)."""
    if not current_user.is_admin():
        flash(get_translation('messages.admin_only_consultants'), 'error')
        return redirect(url_for('index'))
    return redirect(url_for('admin_dashboard'))


@app.route('/consultants/add', methods=['POST'])
@admin_required
def add_consultant():
    """Add a new consultant."""
    display_name = request.form.get('display_name', '').strip()
    if not display_name:
        flash(get_translation('messages.consultant_name_required'), 'error')
        return redirect(url_for('admin_dashboard'))

    conn = get_db_connection()
    result = conn.execute(
        'INSERT INTO consultants (display_name) VALUES (?)',
        (display_name,)
    )
    conn.commit()
    consultant_id = result.lastrowid
    conn.close()

    session['consultant_id'] = consultant_id
    flash(get_translation('messages.consultant_added'), 'success')
    return redirect(url_for('consultants_management'))


@app.route('/consultants/switch/<int:consultant_id>')
@admin_required
def switch_consultant(consultant_id):
    """Switch the active consultant."""
    conn = get_db_connection()
    existing = conn.execute(
        'SELECT id FROM consultants WHERE id = ?',
        (consultant_id,)
    ).fetchone()
    conn.close()

    if not existing:
        flash(get_translation('messages.consultant_not_found'), 'error')
        return redirect(url_for('consultants_management'))

    session['consultant_id'] = consultant_id
    return redirect(request.referrer or url_for('consultants_management'))


def parse_import_payload(payload):
    if not isinstance(payload, dict):
        return None, get_translation('messages.import_invalid_json')

    return {
        'consultant': payload.get('consultant', {}) or {},
        'personal_info': payload.get('personal_info', None),
        'work_experience': payload.get('work_experience', []) or [],
        'education': payload.get('education', []) or [],
        'skills': payload.get('skills', []) or [],
        'projects': payload.get('projects', []) or [],
        'certifications': payload.get('certifications', []) or []
    }, None


def import_consultant_data(conn, consultant_id, payload):
    # Updating, so cleanup the old data first for this consultant to avoid duplicates and stale data
    conn.execute('DELETE FROM personal_info WHERE consultant_id = ?', (consultant_id,))
    conn.execute('DELETE FROM work_experience WHERE consultant_id = ?', (consultant_id,))
    conn.execute('DELETE FROM education WHERE consultant_id = ?', (consultant_id,))
    conn.execute('DELETE FROM skills WHERE consultant_id = ?', (consultant_id,))
    conn.execute('DELETE FROM projects WHERE consultant_id = ?', (consultant_id,))
    conn.execute('DELETE FROM certifications WHERE consultant_id = ?', (consultant_id,))
    
    # Also clean up junction tables for this consultant's entities
    # Note: ON DELETE CASCADE handles this if the parent rows are deleted, 
    # but since we are re-inserting, we ensure a clean slate.

    def normalize_text(value):
        if value is None:
            return ''
        return str(value).strip().lower()

    personal_info = payload.get('personal_info')
    if isinstance(personal_info, dict):
        conn.execute('''
            INSERT INTO personal_info (
                consultant_id, first_name, last_name, email, phone, address, city,
                state, zip_code, country, linkedin_url, github_url, portfolio_url,
                professional_summary
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            consultant_id,
            personal_info.get('first_name', ''),
            personal_info.get('last_name', ''),
            personal_info.get('email', ''),
            personal_info.get('phone', ''),
            personal_info.get('address', ''),
            personal_info.get('city', ''),
            personal_info.get('state', ''),
            personal_info.get('zip_code', ''),
            personal_info.get('country', ''),
            personal_info.get('linkedin_url', ''),
            personal_info.get('github_url', ''),
            personal_info.get('portfolio_url', ''),
            personal_info.get('professional_summary', personal_info.get('summary', ''))
        ))

    # Map skill categories and insert skills
    categories_rows = conn.execute('SELECT id, name FROM skill_categories').fetchall()
    cat_map = {c['name'].lower(): c['id'] for c in categories_rows}
    
    # Ensure "Various" category exists for unknown categories
    if 'various' not in cat_map:
        conn.execute('INSERT INTO skill_categories (name) VALUES (?)', ('Various',))
        conn.commit()
        categories_rows = conn.execute('SELECT id, name FROM skill_categories').fetchall()
        cat_map = {c['name'].lower(): c['id'] for c in categories_rows}
    
    # Track inserted skills by name to link to experience/projects
    skill_name_to_id = {}
    seen_skills = set()

    for skill in payload.get('skills', []):
        cat_name = skill.get('category_name', skill.get('category', '')).strip()
        skill_name = skill.get('skill_name', skill.get('name', ''))
        skill_key = (normalize_text(skill_name), normalize_text(cat_name))
        if skill_key in seen_skills:
            continue
        seen_skills.add(skill_key)
        
        # Map category name to existing category
        # If category doesn't exist, use "Various"
        category_id = None
        if cat_name:
            category_lower = cat_name.lower()
            if category_lower in cat_map:
                category_id = cat_map[category_lower]
            else:
                # Unknown category - map to "Various"
                category_id = cat_map.get('various')
        
        cursor = conn.execute('''
            INSERT INTO skills (
                consultant_id, skill_name, category_id
            ) VALUES (?, ?, ?)
        ''', (
            consultant_id,
            skill_name,
            category_id
        ))
        skill_name_to_id[skill_name.lower()] = cursor.lastrowid

    seen_experience = set()

    for exp in payload.get('work_experience', []):
        exp_key = (
            normalize_text(exp.get('company_name', '')),
            normalize_text(exp.get('position_title', exp.get('job_title', ''))),
            normalize_text(exp.get('location', '')),
            normalize_text(exp.get('start_date', '')),
            normalize_text(exp.get('end_date', ''))
        )
        if exp_key in seen_experience:
            continue
        seen_experience.add(exp_key)
        cursor = conn.execute('''
            INSERT INTO work_experience (
                consultant_id, company_name, position_title, location, start_date,
                end_date, is_current, description, achievements, 
                star_situation, star_tasks, star_actions, star_results,
                display_order
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            consultant_id,
            exp.get('company_name', ''),
            exp.get('position_title', exp.get('job_title', '')),
            exp.get('location', ''),
            exp.get('start_date', ''),
            exp.get('end_date', None),
            1 if exp.get('is_current') or exp.get('end_date', None) is None else 0,
            exp.get('description', ''),
            exp.get('achievements', ''),
            exp.get('star_situation', ''),
            exp.get('star_tasks', ''),
            exp.get('star_actions', ''),
            exp.get('star_results', ''),
            exp.get('display_order', 0)
        ))
        
        # Link skills to work experience
        exp_id = cursor.lastrowid
        linked_skill_ids = set()
        for skill_name in exp.get('skills', []):
            skill_id = skill_name_to_id.get(str(skill_name).strip().lower())
            if skill_id:
                if skill_id in linked_skill_ids:
                    continue
                linked_skill_ids.add(skill_id)
                conn.execute('INSERT INTO experience_skills (experience_id, skill_id) VALUES (?, ?)', (exp_id, skill_id))

    seen_education = set()

    for edu in payload.get('education', []):
        edu_key = (
            normalize_text(edu.get('institution_name', edu.get('institution', ''))),
            normalize_text(edu.get('degree', '')),
            normalize_text(edu.get('field_of_study', '')),
            normalize_text(edu.get('start_date', '')),
            normalize_text(edu.get('end_date', ''))
        )
        if edu_key in seen_education:
            continue
        seen_education.add(edu_key)
        conn.execute('''
            INSERT INTO education (
                consultant_id, institution_name, degree, field_of_study, location,
                start_date, end_date, gpa, honors, description, display_order
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            consultant_id,
            edu.get('institution_name', edu.get('institution', '')),
            edu.get('degree', ''),
            edu.get('field_of_study', ''),
            edu.get('location', ''),
            edu.get('start_date', None),
            edu.get('end_date', None),
            edu.get('gpa', ''),
            edu.get('honors', ''),
            edu.get('description', ''),
            edu.get('display_order', 0)
        ))

    seen_projects = set()

    for project in payload.get('projects', []):
        project_key = (
            normalize_text(project.get('project_name', '')),
            normalize_text(project.get('role', '')),
            normalize_text(project.get('start_date', '')),
            normalize_text(project.get('end_date', ''))
        )
        if project_key in seen_projects:
            continue
        seen_projects.add(project_key)
        cursor = conn.execute('''
            INSERT INTO projects (
                consultant_id, project_name, description,
                start_date, end_date, project_url, github_url, role, achievements,
                display_order
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            consultant_id,
            project.get('project_name', ''),
            project.get('description', ''),
            project.get('start_date', None),
            project.get('end_date', None),
            project.get('project_url', ''),
            project.get('github_url', ''),
            project.get('role', ''),
            project.get('achievements', ''),
            project.get('display_order', 0)
        ))
        
        # Link skills to projects
        proj_id = cursor.lastrowid
        linked_skill_ids = set()
        for skill_name in project.get('skills', []):
            skill_id = skill_name_to_id.get(str(skill_name).strip().lower())
            if skill_id:
                if skill_id in linked_skill_ids:
                    continue
                linked_skill_ids.add(skill_id)
                conn.execute('INSERT INTO project_skills (project_id, skill_id) VALUES (?, ?)', (proj_id, skill_id))

    seen_certifications = set()

    for cert in payload.get('certifications', []):
        cert_key = (
            normalize_text(cert.get('certification_name', '')),
            normalize_text(cert.get('issuing_organization', '')),
            normalize_text(cert.get('issue_date', '')),
            normalize_text(cert.get('credential_id', ''))
        )
        if cert_key in seen_certifications:
            continue
        seen_certifications.add(cert_key)
        conn.execute('''
            INSERT INTO certifications (
                consultant_id, certification_name, issuing_organization, issue_date,
                expiration_date, credential_id, credential_url, description,
                display_order
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            consultant_id,
            cert.get('certification_name', ''),
            cert.get('issuing_organization', ''),
            cert.get('issue_date', None),
            cert.get('expiration_date', None),
            cert.get('credential_id', ''),
            cert.get('credential_url', ''),
            cert.get('description', ''),
            cert.get('display_order', 0)
        ))


@app.route('/consultants/import', methods=['GET', 'POST'])
@admin_required
def import_consultant():
    """Import consultant data from JSON file or textarea."""
    conn = get_db_connection()
    consultants = get_consultants(conn)
    current_consultant_id = resolve_current_consultant_id(conn)
    current_consultant = None
    if current_consultant_id:
        current_consultant = conn.execute(
            'SELECT * FROM consultants WHERE id = ?',
            (current_consultant_id,)
        ).fetchone()
    prefilled_json = None

    if request.method == 'GET':
        # Check if we're prefilling from parsed CV
        cv_id = session.get('parsed_cv_id')
        if cv_id:
            parsed_data, _ = retrieve_parsed_cv_data(cv_id)
            if parsed_data:
                prefilled_json = json.dumps(parsed_data, indent=2)

    if request.method == 'POST':
        # Try to get JSON from textarea first, then from file
        json_textarea = request.form.get('json_textarea', '').strip()
        json_file = request.files.get('json_file')
        import_mode = request.form.get('import_mode')
        
        payload = None
        
        if json_textarea:
            # Parse JSON from textarea
            try:
                payload = json.loads(json_textarea)
            except json.JSONDecodeError as e:
                conn.close()
                flash(f"{get_translation('messages.import_invalid_json')}: {str(e)}", 'error')
                return redirect(url_for('import_consultant'))
        elif json_file:
            # Parse JSON from uploaded file
            try:
                payload = json.load(json_file)
            except json.JSONDecodeError:
                conn.close()
                flash(get_translation('messages.import_invalid_json'), 'error')
                return redirect(url_for('import_consultant'))
        else:
            conn.close()
            flash(get_translation('messages.import_file_required'), 'error')
            return redirect(url_for('import_consultant'))

        parsed_payload, error = parse_import_payload(payload)
        if error:
            conn.close()
            flash(error, 'error')
            return redirect(url_for('import_consultant'))

        if import_mode == 'existing':
            consultant_id = request.form.get('consultant_id')
            if not consultant_id:
                conn.close()
                flash(get_translation('messages.consultant_required'), 'error')
                return redirect(url_for('import_consultant'))

            consultant_id = int(consultant_id)
            exists = conn.execute(
                'SELECT id FROM consultants WHERE id = ?',
                (consultant_id,)
            ).fetchone()
            if not exists:
                conn.close()
                flash(get_translation('messages.consultant_not_found'), 'error')
                return redirect(url_for('import_consultant'))
        else:
            display_name = request.form.get('display_name', '').strip()
            if not display_name:
                display_name = parsed_payload.get('consultant', {}).get('display_name', '').strip()

            if not display_name:
                personal_info = parsed_payload.get('personal_info') or {}
                display_name = (
                    f"{personal_info.get('first_name', '').strip()} "
                    f"{personal_info.get('last_name', '').strip()}"
                ).strip()

            if not display_name:
                display_name = get_translation('consultants.unnamed')

            result = conn.execute(
                'INSERT INTO consultants (display_name) VALUES (?)',
                (display_name,)
            )
            consultant_id = result.lastrowid

        import_consultant_data(conn, consultant_id, parsed_payload)
        conn.commit()
        conn.close()

        session['consultant_id'] = consultant_id
        # Clear parsed CV data if it exists
        cv_id = session.pop('parsed_cv_id', None)
        if cv_id:
            delete_parsed_cv_data(cv_id)
        
        flash(get_translation('messages.import_success'), 'success')
        return redirect(url_for('view_personal_info'))

    conn.close()
    return render_template('import_consultant.html', consultants=consultants, current_consultant=current_consultant, prefilled_json=prefilled_json)


@app.route('/consultants/parse-cv', methods=['GET', 'POST'])
@admin_required
def parse_cv():
    """Parse CV using AI and import consultant data."""
    if request.method == 'POST':
        cv_file = request.files.get('cv_file')
        consultant_name = request.form.get('consultant_name', '').strip()
        ai_provider = request.form.get('ai_provider', 'gemini')
        
        if not cv_file:
            flash(get_translation('messages.cv_file_required'), 'error')
            return redirect(url_for('parse_cv'))
        
        # Extract text from the uploaded file
        cv_text = extract_text_from_file(cv_file)
        if cv_text.startswith('Error') or cv_text.startswith('Unsupported'):
            flash(cv_text, 'error')
            return redirect(url_for('parse_cv'))
        
        # Parse CV with AI
        parsed_data, error = parse_cv_with_ai(cv_text, ai_provider)
        if error:
            flash(f"{get_translation('messages.ai_parse_error')}: {error}", 'error')
            return redirect(url_for('parse_cv'))
        
        if not parsed_data:
            flash(get_translation('messages.no_data_extracted'), 'error')
            return redirect(url_for('parse_cv'))
        
        # Use provided name if specified, otherwise use AI extracted name
        if consultant_name:
            if 'consultant' not in parsed_data:
                parsed_data['consultant'] = {}
            parsed_data['consultant']['display_name'] = consultant_name
        
        # Store parsed data in temporary file and save ID in session
        cv_text_preview = cv_text[:500] + "..." if len(cv_text) > 500 else cv_text
        cv_id = store_parsed_cv_data(parsed_data, cv_text_preview)
        session['parsed_cv_id'] = cv_id
        
        return redirect(url_for('review_parsed_cv'))
    
    return render_template('parse_cv.html', has_groq=bool(GROQ_API_KEY))


@app.route('/consultants/review-parsed-cv')
@admin_required
def review_parsed_cv():
    """Review AI-parsed CV data before importing."""
    cv_id = session.get('parsed_cv_id')
    
    if not cv_id:
        flash(get_translation('messages.no_parsed_data'), 'error')
        return redirect(url_for('parse_cv'))
    
    parsed_data, cv_text_preview = retrieve_parsed_cv_data(cv_id)
    
    if not parsed_data:
        flash(get_translation('messages.no_parsed_data'), 'error')
        return redirect(url_for('parse_cv'))
    
    return render_template('review_parsed_cv.html', 
                         parsed_data=parsed_data, 
                         cv_text_preview=cv_text_preview)


@app.route('/consultants/import-parsed-cv', methods=['POST'])
@admin_required
def import_parsed_cv():
    """Redirect to import screen with AI-parsed CV data prefilled for review/editing."""    
    cv_id = session.get('parsed_cv_id')
    
    if not cv_id:
        flash(get_translation('messages.no_parsed_data'), 'error')
        return redirect(url_for('parse_cv'))
    
    # Verify the data exists
    parsed_data, _ = retrieve_parsed_cv_data(cv_id)
    if not parsed_data:
        flash(get_translation('messages.no_parsed_data'), 'error')
        return redirect(url_for('parse_cv'))
    
    # Redirect to import screen with CV ID in session (will be prefilled in textarea)
    return redirect(url_for('import_consultant'))


@app.route('/assignment-match', methods=['GET', 'POST'])
@login_required
def assignment_match():
    """Match an assignment to one or more consultants using AI."""
    conn = get_db_connection()

    all_consultants = []
    selected_consultant_ids = []
    selected_consultant_names = []
    assignment_description = ''
    ai_provider = 'gemini'
    match_result = None
    match_summary = None
    match_ranking = []
    recommended_names = []
    match_notes = []
    match_result_parse_error = False
    consultants_json_input = None

    if current_user.is_admin():
        all_consultants = conn.execute(
            'SELECT id, display_name FROM consultants ORDER BY display_name, id'
        ).fetchall()
    else:
        if current_user.consultant_id:
            selected_consultant_ids = [current_user.consultant_id]
            own_consultant = conn.execute(
                'SELECT display_name FROM consultants WHERE id = ?',
                (current_user.consultant_id,)
            ).fetchone()
            if own_consultant:
                selected_consultant_names = [own_consultant['display_name']]

    if request.method == 'POST':
        assignment_description = (request.form.get('assignment_description') or '').strip()
        ai_provider = (request.form.get('ai_provider') or 'gemini').strip().lower()

        if not assignment_description:
            conn.close()
            flash(get_translation('messages.assignment_description_required'), 'error')
            return redirect(url_for('assignment_match'))

        if current_user.is_admin():
            raw_ids = request.form.getlist('consultant_ids')
            selected_consultant_ids = []
            for raw_id in raw_ids:
                try:
                    consultant_id = int(raw_id)
                except (TypeError, ValueError):
                    continue
                if consultant_id not in selected_consultant_ids:
                    selected_consultant_ids.append(consultant_id)

            if not selected_consultant_ids:
                conn.close()
                flash(get_translation('messages.assignment_select_consultant_required'), 'error')
                return redirect(url_for('assignment_match'))
        else:
            if not current_user.consultant_id:
                conn.close()
                flash(get_translation('messages.user_no_consultant_access'), 'error')
                return redirect(url_for('index'))
            selected_consultant_ids = [current_user.consultant_id]

        consultant_payloads = []
        for consultant_id in selected_consultant_ids:
            payload = build_consultant_ai_payload(conn, consultant_id)
            if not payload:
                continue
            consultant_payloads.append(payload)
            selected_consultant_names.append(payload['consultant']['display_name'])

        if not consultant_payloads:
            conn.close()
            flash(get_translation('messages.assignment_no_consultant_data'), 'error')
            return redirect(url_for('assignment_match'))

        match_result, consultants_json_input, error = match_assignment_with_ai(
            assignment_description,
            consultant_payloads,
            ai_provider
        )

        if error:
            conn.close()
            flash(f"{get_translation('messages.assignment_match_failed')}: {error}", 'error')
            return redirect(url_for('assignment_match'))

        if match_result:
            match_json_text = match_result.strip()
            if match_json_text.startswith('```json'):
                match_json_text = match_json_text[7:]
            elif match_json_text.startswith('```'):
                match_json_text = match_json_text[3:]

            if match_json_text.endswith('```'):
                match_json_text = match_json_text[:-3]

            match_json_text = match_json_text.strip()

            try:
                parsed_result = json.loads(match_json_text)
                if isinstance(parsed_result, dict):
                    match_summary = parsed_result.get('summary')
                    ranking_data = parsed_result.get('ranking')
                    if isinstance(ranking_data, list):
                        match_ranking = ranking_data
                    recommended_data = parsed_result.get('recommended_consultant_names')
                    if isinstance(recommended_data, list):
                        recommended_names = recommended_data
                    notes_data = parsed_result.get('notes')
                    if isinstance(notes_data, list):
                        match_notes = notes_data
                else:
                    match_result_parse_error = True
            except json.JSONDecodeError:
                match_result_parse_error = True

        flash(get_translation('messages.assignment_match_success'), 'success')

    conn.close()
    return render_template(
        'assignment_match.html',
        all_consultants=all_consultants,
        selected_consultant_ids=selected_consultant_ids,
        selected_consultant_names=selected_consultant_names,
        assignment_description=assignment_description,
        ai_provider=ai_provider,
        match_result=match_result,
        match_summary=match_summary,
        match_ranking=match_ranking,
        recommended_names=recommended_names,
        match_notes=match_notes,
        match_result_parse_error=match_result_parse_error,
        consultants_json_input=consultants_json_input,
        has_groq=bool(GROQ_API_KEY)
    )


@app.route('/consultants/delete/<int:consultant_id>', methods=['POST'])
@admin_required
def delete_consultant(consultant_id):
    """Delete a consultant."""
    conn = get_db_connection()
    
    # Check if consultant exists
    consultant = conn.execute(
        'SELECT id, display_name FROM consultants WHERE id = ?',
        (consultant_id,)
    ).fetchone()
    
    if not consultant:
        conn.close()
        flash(get_translation('messages.consultant_not_found'), 'error')
        return redirect(url_for('consultants_management'))
    
    # Prevent deleting if this is the only consultant
    consultant_count = conn.execute('SELECT COUNT(*) FROM consultants').fetchone()[0]
    if consultant_count <= 1:
        conn.close()
        flash(get_translation('messages.consultant_delete_last'), 'error')
        return redirect(url_for('consultants_management'))
    
    # If deleting current consultant, switch to another one
    current_consultant_id = session.get('consultant_id')
    if current_consultant_id == consultant_id:
        # Find another consultant to switch to
        other_consultant = conn.execute(
            'SELECT id FROM consultants WHERE id != ? ORDER BY id LIMIT 1',
            (consultant_id,)
        ).fetchone()
        if other_consultant:
            session['consultant_id'] = other_consultant['id']
        else:
            # This should not happen due to the count check above
            session.pop('consultant_id', None)
    
    # Delete all related data
    conn.execute('DELETE FROM users WHERE consultant_id = ?', (consultant_id,))
    conn.execute('DELETE FROM personal_info WHERE consultant_id = ?', (consultant_id,))
    conn.execute('DELETE FROM work_experience WHERE consultant_id = ?', (consultant_id,))
    conn.execute('DELETE FROM education WHERE consultant_id = ?', (consultant_id,))
    conn.execute('DELETE FROM skills WHERE consultant_id = ?', (consultant_id,))
    conn.execute('DELETE FROM projects WHERE consultant_id = ?', (consultant_id,))
    conn.execute('DELETE FROM certifications WHERE consultant_id = ?', (consultant_id,))
    
    # Delete the consultant
    conn.execute('DELETE FROM consultants WHERE id = ?', (consultant_id,))
    conn.commit()
    conn.close()
    
    flash(get_translation('messages.consultant_deleted'), 'success')
    return redirect(url_for('consultants_management'))


@app.route('/consultants/export/<int:consultant_id>')
@admin_required
def export_consultant(consultant_id):
    """Export a consultants data as JSON."""
    conn = get_db_connection()
    payload = build_consultant_ai_payload(conn, consultant_id)

    if not payload:
        conn.close()
        flash(get_translation('messages.consultant_not_found'), 'error')
        return redirect(url_for('admin_dashboard'))

    payload['version'] = 1

    conn.close()

    # Sanitize display name for filename
    safe_name = "".join([c if c.isalnum() else "_" for c in payload['consultant']['display_name']])
    filename = f"{safe_name}.json"
    
    response = Response(
        json.dumps(payload, indent=2),
        mimetype='application/json'
    )
    response.headers['Content-Disposition'] = f'attachment; filename={filename}'
    return response


# ==================== Personal Info Routes ====================

@app.route('/personal-info')
@login_required
def view_personal_info():
    """View personal information."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    personal_info = conn.execute(
        'SELECT * FROM personal_info WHERE consultant_id = ? ORDER BY id DESC LIMIT 1',
        (consultant_id,)
    ).fetchone()
    conn.close()
    return render_template('personal_info.html', info=personal_info)


@app.route('/personal-info/edit', methods=['GET', 'POST'])
@login_required
def edit_personal_info():
    """Edit or create personal information."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    
    if request.method == 'POST':
        # Get form data
        data = {
            'first_name': request.form['first_name'],
            'last_name': request.form['last_name'],
            'email': request.form['email'],
            'phone': request.form.get('phone', ''),
            'address': request.form.get('address', ''),
            'city': request.form.get('city', ''),
            'state': request.form.get('state', ''),
            'zip_code': request.form.get('zip_code', ''),
            'country': request.form.get('country', ''),
            'linkedin_url': request.form.get('linkedin_url', ''),
            'github_url': request.form.get('github_url', ''),
            'portfolio_url': request.form.get('portfolio_url', ''),
            'professional_summary': request.form.get('professional_summary', '')
        }
        
        # Check if record exists
        existing = conn.execute(
            'SELECT id FROM personal_info WHERE consultant_id = ? LIMIT 1',
            (consultant_id,)
        ).fetchone()
        
        if existing:
            # Update existing record
            conn.execute('''
                UPDATE personal_info SET
                    first_name = ?, last_name = ?, email = ?, phone = ?,
                    address = ?, city = ?, state = ?, zip_code = ?, country = ?,
                    linkedin_url = ?, github_url = ?, portfolio_url = ?,
                    professional_summary = ?, updated_at = CURRENT_TIMESTAMP
                WHERE id = ? AND consultant_id = ?
            ''', (data['first_name'], data['last_name'], data['email'], data['phone'],
                  data['address'], data['city'], data['state'], data['zip_code'], data['country'],
                  data['linkedin_url'], data['github_url'], data['portfolio_url'],
                  data['professional_summary'], existing['id'], consultant_id))
            flash(get_translation('messages.personal_info_updated'), 'success')
        else:
            # Insert new record
            conn.execute('''
                INSERT INTO personal_info (
                    consultant_id, first_name, last_name, email, phone, address, city, state,
                    zip_code, country, linkedin_url, github_url, portfolio_url, professional_summary
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                consultant_id,
                data['first_name'], data['last_name'], data['email'], data['phone'],
                data['address'], data['city'], data['state'], data['zip_code'], data['country'],
                data['linkedin_url'], data['github_url'], data['portfolio_url'],
                data['professional_summary']
            ))
            flash(get_translation('messages.personal_info_saved'), 'success')

        display_name = f"{data['first_name']} {data['last_name']}".strip()
        if display_name:
            conn.execute(
                'UPDATE consultants SET display_name = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?',
                (display_name, consultant_id)
            )
        
        conn.commit()
        conn.close()
        return redirect(url_for('view_personal_info'))
    
    # GET request - show form
    personal_info = conn.execute(
        'SELECT * FROM personal_info WHERE consultant_id = ? ORDER BY id DESC LIMIT 1',
        (consultant_id,)
    ).fetchone()
    conn.close()
    return render_template('edit_personal_info.html', info=personal_info)


# ==================== Work Experience Routes ====================

@app.route('/work-experience')
@login_required
def view_work_experience():
    """View all work experience entries."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    experiences_raw = conn.execute(
        'SELECT * FROM work_experience WHERE consultant_id = ? ORDER BY display_order, start_date DESC',
        (consultant_id,)
    ).fetchall()
    
    experiences = []
    for exp in experiences_raw:
        exp_dict = dict(exp)
        # Fetch skills for this experience
        skills = conn.execute('''
            SELECT s.skill_name 
            FROM skills s
            JOIN experience_skills es ON s.id = es.skill_id
            WHERE es.experience_id = ?
            ORDER BY s.category, s.skill_name
        ''', (exp['id'],)).fetchall()
        exp_dict['skills'] = [s['skill_name'] for s in skills]
        experiences.append(exp_dict)
        
    conn.close()
    return render_template('work_experience.html', experiences=experiences)


@app.route('/work-experience/add', methods=['GET', 'POST'])
@login_required
def add_work_experience():
    """Add new work experience."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    
    if request.method == 'POST':
        cursor = conn.execute('''
            INSERT INTO work_experience (
                consultant_id, company_name, position_title, location, start_date, end_date,
                is_current, description, achievements, star_situation, star_tasks, 
                star_actions, star_results, display_order
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            consultant_id,
            request.form['company_name'],
            request.form['position_title'],
            request.form.get('location', ''),
            request.form['start_date'],
            request.form.get('end_date', None) if not request.form.get('is_current') else None,
            1 if request.form.get('is_current') else 0,
            request.form.get('description', ''),
            request.form.get('achievements', ''),
            request.form.get('star_situation', ''),
            request.form.get('star_tasks', ''),
            request.form.get('star_actions', ''),
            request.form.get('star_results', ''),
            request.form.get('display_order', 0)
        ))
        
        experience_id = cursor.lastrowid
        
        # Handle linked skills
        skill_ids = request.form.getlist('skill_ids')
        for skill_id in skill_ids:
            conn.execute(
                'INSERT INTO experience_skills (experience_id, skill_id) VALUES (?, ?)',
                (experience_id, skill_id)
            )
        
        conn.commit()
        conn.close()
        flash(get_translation('messages.work_experience_added'), 'success')
        return redirect(url_for('view_work_experience'))
    
    skills = conn.execute('''
        SELECT s.id, s.skill_name, c.name as category_name
        FROM skills s
        LEFT JOIN skill_categories c ON s.category_id = c.id
        WHERE s.consultant_id = ?
        ORDER BY c.name, s.skill_name
    ''', (consultant_id,)).fetchall()
    conn.close()
    return render_template('edit_work_experience.html', experience=None, all_skills=skills)


@app.route('/work-experience/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_work_experience(id):
    """Edit existing work experience."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    
    if request.method == 'POST':
        conn.execute('''
            UPDATE work_experience SET
                company_name = ?, position_title = ?, location = ?, start_date = ?,
                end_date = ?, is_current = ?, description = ?, achievements = ?,
                star_situation = ?, star_tasks = ?, star_actions = ?, star_results = ?,
                display_order = ?, updated_at = CURRENT_TIMESTAMP
            WHERE id = ? AND consultant_id = ?
        ''', (
            request.form['company_name'],
            request.form['position_title'],
            request.form.get('location', ''),
            request.form['start_date'],
            request.form.get('end_date', None) if not request.form.get('is_current') else None,
            1 if request.form.get('is_current') else 0,
            request.form.get('description', ''),
            request.form.get('achievements', ''),
            request.form.get('star_situation', ''),
            request.form.get('star_tasks', ''),
            request.form.get('star_actions', ''),
            request.form.get('star_results', ''),
            request.form.get('display_order', 0),
            id,
            consultant_id
        ))
        
        # Update linked skills
        conn.execute('DELETE FROM experience_skills WHERE experience_id = ?', (id,))
        skill_ids = request.form.getlist('skill_ids')
        for skill_id in skill_ids:
            conn.execute(
                'INSERT INTO experience_skills (experience_id, skill_id) VALUES (?, ?)',
                (id, skill_id)
            )
        
        conn.commit()
        conn.close()
        flash(get_translation('messages.work_experience_updated'), 'success')
        return redirect(url_for('view_work_experience'))
    
    experience = conn.execute(
        'SELECT * FROM work_experience WHERE id = ? AND consultant_id = ?',
        (id, consultant_id)
    ).fetchone()
    
    if not experience:
        conn.close()
        flash(get_translation('messages.work_experience_not_found'), 'error')
        return redirect(url_for('view_work_experience'))
        
    all_skills = conn.execute('''
        SELECT s.id, s.skill_name, c.name as category_name
        FROM skills s
        LEFT JOIN skill_categories c ON s.category_id = c.id
        WHERE s.consultant_id = ?
        ORDER BY c.name, s.skill_name
    ''', (consultant_id,)).fetchall()
    
    current_skill_ids = [row['skill_id'] for row in conn.execute(
        'SELECT skill_id FROM experience_skills WHERE experience_id = ?',
        (id,)
    ).fetchall()]
    
    conn.close()
    return render_template('edit_work_experience.html', 
                         experience=experience, 
                         all_skills=all_skills,
                         current_skill_ids=current_skill_ids)


@app.route('/work-experience/delete/<int:id>', methods=['POST'])
@login_required
def delete_work_experience(id):
    """Delete work experience."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    conn.execute(
        'DELETE FROM work_experience WHERE id = ? AND consultant_id = ?',
        (id, consultant_id)
    )
    conn.commit()
    conn.close()
    flash(get_translation('messages.work_experience_deleted'), 'success')
    return redirect(url_for('view_work_experience'))


# ==================== Education Routes ====================

@app.route('/education')
@login_required
def view_education():
    """View all education entries."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    education = conn.execute(
        'SELECT * FROM education WHERE consultant_id = ? ORDER BY display_order, end_date DESC',
        (consultant_id,)
    ).fetchall()
    conn.close()
    return render_template('education.html', education=education)


@app.route('/education/add', methods=['GET', 'POST'])
@login_required
def add_education():
    """Add new education entry."""
    if request.method == 'POST':
        conn = get_db_connection()
        consultant_id = resolve_current_consultant_id(conn)
        
        conn.execute('''
            INSERT INTO education (
                consultant_id, institution_name, degree, field_of_study, location, start_date,
                end_date, gpa, honors, description, display_order
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            consultant_id,
            request.form['institution_name'],
            request.form['degree'],
            request.form.get('field_of_study', ''),
            request.form.get('location', ''),
            request.form.get('start_date', None),
            request.form.get('end_date', None),
            request.form.get('gpa', ''),
            request.form.get('honors', ''),
            request.form.get('description', ''),
            request.form.get('display_order', 0)
        ))
        
        conn.commit()
        conn.close()
        flash(get_translation('messages.education_added'), 'success')
        return redirect(url_for('view_education'))
    
    return render_template('edit_education.html', education=None)


@app.route('/education/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_education(id):
    """Edit existing education entry."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    
    if request.method == 'POST':
        conn.execute('''
            UPDATE education SET
                institution_name = ?, degree = ?, field_of_study = ?, location = ?,
                start_date = ?, end_date = ?, gpa = ?, honors = ?, description = ?,
                display_order = ?, updated_at = CURRENT_TIMESTAMP
            WHERE id = ? AND consultant_id = ?
        ''', (
            request.form['institution_name'],
            request.form['degree'],
            request.form.get('field_of_study', ''),
            request.form.get('location', ''),
            request.form.get('start_date', None),
            request.form.get('end_date', None),
            request.form.get('gpa', ''),
            request.form.get('honors', ''),
            request.form.get('description', ''),
            request.form.get('display_order', 0),
            id,
            consultant_id
        ))
        
        conn.commit()
        conn.close()
        flash(get_translation('messages.education_updated'), 'success')
        return redirect(url_for('view_education'))
    
    education = conn.execute(
        'SELECT * FROM education WHERE id = ? AND consultant_id = ?',
        (id, consultant_id)
    ).fetchone()
    conn.close()
    return render_template('edit_education.html', education=education)


@app.route('/education/delete/<int:id>', methods=['POST'])
@login_required
def delete_education(id):
    """Delete education entry."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    conn.execute(
        'DELETE FROM education WHERE id = ? AND consultant_id = ?',
        (id, consultant_id)
    )
    conn.commit()
    conn.close()
    flash(get_translation('messages.education_deleted'), 'success')
    return redirect(url_for('view_education'))


# ==================== Skills Routes ====================

@app.route('/skills')
@login_required
def view_skills():
    """View all skills."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    skills = conn.execute('''
        SELECT s.*, c.name as category_name
        FROM skills s
        LEFT JOIN skill_categories c ON s.category_id = c.id
        WHERE s.consultant_id = ?
        ORDER BY c.name, s.skill_name
    ''', (consultant_id,)).fetchall()
    conn.close()
    return render_template('skills.html', skills=skills)


@app.route('/skills/add', methods=['GET', 'POST'])
@login_required
def add_skill():
    """Add new skill."""
    if request.method == 'POST':
        conn = get_db_connection()
        consultant_id = resolve_current_consultant_id(conn)
        
        conn.execute('''
            INSERT INTO skills (
                consultant_id, skill_name, category_id
            ) VALUES (?, ?, ?)
        ''', (
            consultant_id,
            request.form['skill_name'],
            request.form.get('category_id')
        ))
        
        conn.commit()
        conn.close()
        flash(get_translation('messages.skill_added'), 'success')
        return redirect(url_for('view_skills'))
    
    categories = get_skill_categories()
    return render_template('edit_skill.html', skill=None, categories=categories)


@app.route('/skills/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_skill(id):
    """Edit existing skill."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    
    if request.method == 'POST':
        conn.execute('''
            UPDATE skills SET
                skill_name = ?, category_id = ?,
                updated_at = CURRENT_TIMESTAMP
            WHERE id = ? AND consultant_id = ?
        ''', (
            request.form['skill_name'],
            request.form.get('category_id'),
            id,
            consultant_id
        ))
        
        conn.commit()
        conn.close()
        flash(get_translation('messages.skill_updated'), 'success')
        return redirect(url_for('view_skills'))
    
    skill = conn.execute(
        'SELECT * FROM skills WHERE id = ? AND consultant_id = ?',
        (id, consultant_id)
    ).fetchone()
    categories = get_skill_categories()
    conn.close()
    return render_template('edit_skill.html', skill=skill, categories=categories)


@app.route('/skills/delete/<int:id>', methods=['POST'])
@login_required
def delete_skill(id):
    """Delete skill."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    conn.execute(
        'DELETE FROM skills WHERE id = ? AND consultant_id = ?',
        (id, consultant_id)
    )
    conn.commit()
    conn.close()
    flash(get_translation('messages.skill_deleted'), 'success')
    return redirect(url_for('view_skills'))


# ==================== Projects Routes ====================

@app.route('/projects')
@login_required
def view_projects():
    """View all projects."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    projects_raw = conn.execute(
        'SELECT * FROM projects WHERE consultant_id = ? ORDER BY display_order, start_date DESC',
        (consultant_id,)
    ).fetchall()
    
    projects = []
    for proj in projects_raw:
        proj_dict = dict(proj)
        # Fetch skills for this project
        skills = conn.execute('''
            SELECT s.skill_name 
            FROM skills s
            JOIN project_skills ps ON s.id = ps.skill_id
            WHERE ps.project_id = ?
            ORDER BY s.skill_name
        ''', (proj['id'],)).fetchall()
        proj_dict['skills'] = [s['skill_name'] for s in skills]
        projects.append(proj_dict)
        
    conn.close()
    return render_template('projects.html', projects=projects)


@app.route('/projects/add', methods=['GET', 'POST'])
@login_required
def add_project():
    """Add new project."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    
    if request.method == 'POST':
        cursor = conn.execute('''
            INSERT INTO projects (
                consultant_id, project_name, description, start_date, end_date,
                project_url, github_url, role, achievements, display_order
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            consultant_id,
            request.form['project_name'],
            request.form.get('description', ''),
            request.form.get('start_date', None),
            request.form.get('end_date', None),
            request.form.get('project_url', ''),
            request.form.get('github_url', ''),
            request.form.get('role', ''),
            request.form.get('achievements', ''),
            request.form.get('display_order', 0)
        ))
        
        project_id = cursor.lastrowid
        
        # Handle linked skills
        skill_ids = request.form.getlist('skill_ids')
        for skill_id in skill_ids:
            conn.execute(
                'INSERT INTO project_skills (project_id, skill_id) VALUES (?, ?)',
                (project_id, skill_id)
            )
        
        conn.commit()
        conn.close()
        flash(get_translation('messages.project_added'), 'success')
        return redirect(url_for('view_projects'))
    
    all_skills = conn.execute('''
        SELECT s.id, s.skill_name, c.name as category_name
        FROM skills s
        LEFT JOIN skill_categories c ON s.category_id = c.id
        WHERE s.consultant_id = ?
        ORDER BY c.name, s.skill_name
    ''', (consultant_id,)).fetchall()
    conn.close()
    return render_template('edit_project.html', project=None, all_skills=all_skills)


@app.route('/projects/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_project(id):
    """Edit existing project."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    
    if request.method == 'POST':
        conn.execute('''
            UPDATE projects SET
                project_name = ?, description = ?, start_date = ?,
                end_date = ?, project_url = ?, github_url = ?, role = ?, achievements = ?,
                display_order = ?, updated_at = CURRENT_TIMESTAMP
            WHERE id = ? AND consultant_id = ?
        ''', (
            request.form['project_name'],
            request.form.get('description', ''),
            request.form.get('start_date', None),
            request.form.get('end_date', None),
            request.form.get('project_url', ''),
            request.form.get('github_url', ''),
            request.form.get('role', ''),
            request.form.get('achievements', ''),
            request.form.get('display_order', 0),
            id,
            consultant_id
        ))
        
        # Update linked skills
        conn.execute('DELETE FROM project_skills WHERE project_id = ?', (id,))
        skill_ids = request.form.getlist('skill_ids')
        for skill_id in skill_ids:
            conn.execute(
                'INSERT INTO project_skills (project_id, skill_id) VALUES (?, ?)',
                (id, skill_id)
            )
        
        conn.commit()
        conn.close()
        flash(get_translation('messages.project_updated'), 'success')
        return redirect(url_for('view_projects'))
    
    project = conn.execute(
        'SELECT * FROM projects WHERE id = ? AND consultant_id = ?',
        (id, consultant_id)
    ).fetchone()
    
    if not project:
        conn.close()
        flash(get_translation('messages.project_not_found'), 'error')
        return redirect(url_for('view_projects'))
        
    all_skills = conn.execute('''
        SELECT s.id, s.skill_name, c.name as category_name
        FROM skills s
        LEFT JOIN skill_categories c ON s.category_id = c.id
        WHERE s.consultant_id = ?
        ORDER BY c.name, s.skill_name
    ''', (consultant_id,)).fetchall()
    
    current_skill_ids = [row['skill_id'] for row in conn.execute(
        'SELECT skill_id FROM project_skills WHERE project_id = ?',
        (id,)
    ).fetchall()]
    
    conn.close()
    return render_template('edit_project.html', 
                         project=project, 
                         all_skills=all_skills,
                         current_skill_ids=current_skill_ids)


@app.route('/projects/delete/<int:id>', methods=['POST'])
@login_required
def delete_project(id):
    """Delete project."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    conn.execute(
        'DELETE FROM projects WHERE id = ? AND consultant_id = ?',
        (id, consultant_id)
    )
    conn.commit()
    conn.close()
    flash(get_translation('messages.project_deleted'), 'success')
    return redirect(url_for('view_projects'))


# ==================== Certifications Routes ====================

@app.route('/certifications')
@login_required
def view_certifications():
    """View all certifications."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    certifications_raw = conn.execute(
        'SELECT * FROM certifications WHERE consultant_id = ? ORDER BY display_order, issue_date DESC',
        (consultant_id,)
    ).fetchall()
    conn.close()
    
    # Add issue_year to each certification
    certifications = []
    for cert in certifications_raw:
        cert_dict = dict(cert)
        if cert_dict.get('issue_date'):
            issue_date_str = str(cert_dict['issue_date'])
            if issue_date_str and len(issue_date_str) >= 4:
                cert_dict['issue_year'] = issue_date_str[:4]
        certifications.append(cert_dict)
    
    return render_template('certifications.html', certifications=certifications)


@app.route('/certifications/add', methods=['GET', 'POST'])
@login_required
def add_certification():
    """Add new certification."""
    if request.method == 'POST':
        conn = get_db_connection()
        consultant_id = resolve_current_consultant_id(conn)
        
        conn.execute('''
            INSERT INTO certifications (
                consultant_id, certification_name, issuing_organization, issue_date, expiration_date,
                credential_id, credential_url, description, display_order
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            consultant_id,
            request.form['certification_name'],
            request.form['issuing_organization'],
            request.form.get('issue_date', None),
            request.form.get('expiration_date', None),
            request.form.get('credential_id', ''),
            request.form.get('credential_url', ''),
            request.form.get('description', ''),
            request.form.get('display_order', 0)
        ))
        
        conn.commit()
        conn.close()
        flash(get_translation('messages.certification_added'), 'success')
        return redirect(url_for('view_certifications'))
    
    return render_template('edit_certification.html', certification=None)


@app.route('/certifications/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_certification(id):
    """Edit existing certification."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    
    if request.method == 'POST':
        conn.execute('''
            UPDATE certifications SET
                certification_name = ?, issuing_organization = ?, issue_date = ?,
                expiration_date = ?, credential_id = ?, credential_url = ?, description = ?,
                display_order = ?, updated_at = CURRENT_TIMESTAMP
            WHERE id = ? AND consultant_id = ?
        ''', (
            request.form['certification_name'],
            request.form['issuing_organization'],
            request.form.get('issue_date', None),
            request.form.get('expiration_date', None),
            request.form.get('credential_id', ''),
            request.form.get('credential_url', ''),
            request.form.get('description', ''),
            request.form.get('display_order', 0),
            id,
            consultant_id
        ))
        
        conn.commit()
        conn.close()
        flash(get_translation('messages.certification_updated'), 'success')
        return redirect(url_for('view_certifications'))
    
    certification = conn.execute(
        'SELECT * FROM certifications WHERE id = ? AND consultant_id = ?',
        (id, consultant_id)
    ).fetchone()
    conn.close()
    return render_template('edit_certification.html', certification=certification)


@app.route('/certifications/delete/<int:id>', methods=['POST'])
@login_required
def delete_certification(id):
    """Delete certification."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    conn.execute(
        'DELETE FROM certifications WHERE id = ? AND consultant_id = ?',
        (id, consultant_id)
    )
    conn.commit()
    conn.close()
    flash(get_translation('messages.certification_deleted'), 'success')
    return redirect(url_for('view_certifications'))


# ==================== CV Export Routes ====================

@app.route('/export-cv')
@login_required
def export_cv():
    """Display CV export page."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    
    # Check if there's any data to export
    personal_info = conn.execute(
        'SELECT * FROM personal_info WHERE consultant_id = ?',
        (consultant_id,)
    ).fetchone()
    
    work_count = conn.execute(
        'SELECT COUNT(*) as count FROM work_experience WHERE consultant_id = ?',
        (consultant_id,)
    ).fetchone()['count']
    
    conn.close()
    
    has_data = personal_info is not None or work_count > 0
    
    return render_template('export_cv.html', has_data=has_data)


@app.route('/export-cv/preview')
@login_required
def preview_cv():
    """Preview CV in HTML format."""
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    
    # Get all CV data
    cv_data = get_cv_data(conn, consultant_id)
    conn.close()
    
    return render_template('cv_template.html', **cv_data)


@app.route('/export-cv/download', methods=['POST'])
@login_required
def export_cv_download():
    """Download CV in selected format."""
    format_type = request.form.get('format', 'html')
    
    conn = get_db_connection()
    consultant_id = resolve_current_consultant_id(conn)
    
    # Get all CV data
    cv_data = get_cv_data(conn, consultant_id)
    conn.close()
    
    # Get consultant name for filename
    if cv_data['personal_info']:
        filename = f"CV_{cv_data['personal_info']['first_name']}_{cv_data['personal_info']['last_name']}"
    else:
        filename = "CV_Export"
    
    if format_type == 'html':
        html_content = render_template('cv_template.html', **cv_data)
        return Response(
            html_content,
            mimetype='text/html',
            headers={'Content-Disposition': f'attachment;filename={filename}.html'}
        )
    
    elif format_type == 'pdf':
        try:
            import pdfkit
            html_content = render_template('cv_template.html', **cv_data)
            pdf = pdfkit.from_string(html_content, False)
            return Response(
                pdf,
                mimetype='application/pdf',
                headers={'Content-Disposition': f'attachment;filename={filename}.pdf'}
            )
        except Exception as e:
            flash(f'PDF generation failed: {str(e)}. Please install wkhtmltopdf or use HTML export.', 'error')
            return redirect(url_for('export_cv'))
    
    elif format_type == 'docx':
        try:
            # Generate HTML from template
            html_content = render_template('cv_template.html', **cv_data)
            
            # Convert HTML to DOCX
            from html2docx import html_to_docx
            from io import BytesIO
            
            doc_io = BytesIO()
            html_to_docx(html_content, doc_file=doc_io)
            doc_io.seek(0)
            
            return Response(
                doc_io.getvalue(),
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={'Content-Disposition': f'attachment;filename={filename}.docx'}
            )
            
        except ImportError:
            # Fallback: If html2docx is not available, use basic python-docx conversion with BeautifulSoup
            try:
                from docx import Document
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from html.parser import HTMLParser
                from io import BytesIO
                import re
                
                # Generate HTML from template
                html_content = render_template('cv_template.html', **cv_data)
                
                # Simple HTML to docx conversion
                doc = Document()
                
                # Parse HTML and extract text content
                class HTMLExtractor(HTMLParser):
                    def __init__(self):
                        super().__init__()
                        self.in_tag = None
                        self.current_text = []
                        self.current_heading_level = None
                        
                    def handle_starttag(self, tag, attrs):
                        if tag in ['h1', 'h2', 'h3']:
                            self.in_tag = tag
                    
                    def handle_endtag(self, tag):
                        if tag in ['h1', 'h2', 'h3']:
                            self.in_tag = None
                    
                    def handle_data(self, data):
                        if data.strip():
                            self.current_text.append(data.strip())
                
                # Extract text content and structure from HTML
                import re
                # Remove style tags
                html_no_styles = re.sub(r'<style>.*?</style>', '', html_content, flags=re.DOTALL)
                # Remove script tags
                html_no_scripts = re.sub(r'<script>.*?</script>', '', html_no_styles, flags=re.DOTALL)
                
                # Simple heading detection
                headings = re.findall(r'<h([1-3]).*?>(.*?)</h\1>', html_no_scripts, re.DOTALL)
                paragraphs = re.findall(r'<p.*?>(.*?)</p>', html_no_scripts, re.DOTALL)
                
                # Add headings to document
                for level, content in headings:
                    level_map = {'1': 0, '2': 1, '3': 2}
                    clean_content = re.sub(r'<[^>]+>', '', content).strip()
                    if clean_content:
                        doc.add_heading(clean_content, int(level_map.get(level, 1)))
                
                # Add paragraphs
                for content in paragraphs:
                    clean_content = re.sub(r'<[^>]+>', '', content).strip()
                    if clean_content:
                        doc.add_paragraph(clean_content)
                
                doc_io = BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                return Response(
                    doc_io.getvalue(),
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    headers={'Content-Disposition': f'attachment;filename={filename}.docx'}
                )
                
            except Exception as e:
                flash(f'Word document generation failed: {str(e)}. Please ensure html2docx is installed: pip install html2docx', 'error')
                return redirect(url_for('export_cv'))
        
        except Exception as e:
            flash(f'Word document generation failed: {str(e)}', 'error')
            return redirect(url_for('export_cv'))
    
    elif format_type == 'json':
        # Export as JSON (similar to export_consultant but for current consultant)
        conn = get_db_connection()
        consultant_id = resolve_current_consultant_id(conn)
        
        consultant = conn.execute(
            'SELECT id, display_name FROM consultants WHERE id = ?',
            (consultant_id,)
        ).fetchone()
        
        payload = {
            'version': 1,
            'consultant': {
                'display_name': consultant['display_name'] if consultant else 'Unknown'
            },
            'personal_info': None,
            'work_experience': [],
            'education': [],
            'skills': [],
            'projects': [],
            'certifications': []
        }

        personal_info = conn.execute(
            'SELECT * FROM personal_info WHERE consultant_id = ? ORDER BY id DESC LIMIT 1',
            (consultant_id,)
        ).fetchone()
        if personal_info:
            payload['personal_info'] = dict(personal_info)
            payload['personal_info'].pop('id', None)
            payload['personal_info'].pop('consultant_id', None)
            payload['personal_info'].pop('created_at', None)
            payload['personal_info'].pop('updated_at', None)

        for table, key in [
            ('work_experience', 'work_experience'),
            ('education', 'education'),
            ('skills', 'skills'),
            ('projects', 'projects'),
            ('certifications', 'certifications')
        ]:
            rows = conn.execute(
                f'SELECT * FROM {table} WHERE consultant_id = ? ORDER BY id',
                (consultant_id,)
            ).fetchall()
            payload[key] = []
            for row in rows:
                row_data = dict(row)
                
                # Special handling for skills to include category name
                if table == 'skills' and row_data.get('category_id'):
                    cat = conn.execute('SELECT name FROM skill_categories WHERE id = ?', (row_data['category_id'],)).fetchone()
                    if cat:
                        row_data['category_name'] = cat['name']

                # Special handling for experience and projects to include linked skills
                if table == 'work_experience':
                    skills = conn.execute('''
                        SELECT s.skill_name 
                        FROM skills s
                        JOIN experience_skills es ON s.id = es.skill_id
                        WHERE es.experience_id = ?
                    ''', (row['id'],)).fetchall()
                    row_data['skills'] = [s['skill_name'] for s in skills]
                
                if table == 'projects':
                    skills = conn.execute('''
                        SELECT s.skill_name 
                        FROM skills s
                        JOIN project_skills ps ON s.id = ps.skill_id
                        WHERE ps.project_id = ?
                    ''', (row['id'],)).fetchall()
                    row_data['skills'] = [s['skill_name'] for s in skills]

                row_data.pop('id', None)
                row_data.pop('consultant_id', None)
                row_data.pop('created_at', None)
                row_data.pop('updated_at', None)
                row_data.pop('category', None)  # Remove old text category field if exists
                row_data.pop('years_of_experience', None)  # Remove legacy field if it exists
                payload[key].append(row_data)

        conn.close()
        
        return Response(
            json.dumps(payload, indent=2),
            mimetype='application/json',
            headers={'Content-Disposition': f'attachment;filename={filename}.json'}
        )
    
    flash('Invalid format selected', 'error')
    return redirect(url_for('export_cv'))


def get_cv_data(conn, consultant_id):
    """Helper function to retrieve all CV data for a consultant."""
    
    # Personal Info
    personal_info = conn.execute(
        'SELECT * FROM personal_info WHERE consultant_id = ?',
        (consultant_id,)
    ).fetchone()
    
    # Work Experience
    work_experience_rows = conn.execute(
        'SELECT * FROM work_experience WHERE consultant_id = ? ORDER BY display_order, start_date DESC',
        (consultant_id,)
    ).fetchall()

    work_experiences = []
    for row in work_experience_rows:
        exp_data = dict(row)
        skills = conn.execute('''
            SELECT s.skill_name
            FROM skills s
            JOIN experience_skills es ON s.id = es.skill_id
            WHERE es.experience_id = ?
            ORDER BY s.skill_name
        ''', (row['id'],)).fetchall()
        exp_data['skills'] = [s['skill_name'] for s in skills]
        work_experiences.append(exp_data)
    
    # Education
    education = conn.execute(
        'SELECT * FROM education WHERE consultant_id = ? ORDER BY display_order, start_date DESC',
        (consultant_id,)
    ).fetchall()
    
    # Certifications
    certifications_raw = conn.execute(
        'SELECT * FROM certifications WHERE consultant_id = ? ORDER BY display_order, issue_date DESC',
        (consultant_id,)
    ).fetchall()
    
    # Add issue_year to each certification
    certifications = []
    for cert in certifications_raw:
        cert_dict = dict(cert)
        if cert_dict.get('issue_date'):
            issue_date_str = str(cert_dict['issue_date'])
            if issue_date_str and len(issue_date_str) >= 4:
                cert_dict['issue_year'] = issue_date_str[:4]
        certifications.append(cert_dict)
    
    # Skills grouped by category
    skills = conn.execute('''
        SELECT s.*, c.name as category_name
        FROM skills s
        LEFT JOIN skill_categories c ON s.category_id = c.id
        WHERE s.consultant_id = ?
        ORDER BY c.name, s.skill_name
    ''', (consultant_id,)).fetchall()
    
    skills_by_category = {}
    for skill in skills:
        category = skill['category_name'] or get_translation('skills.uncategorized')
        if category not in skills_by_category:
            skills_by_category[category] = []
        skills_by_category[category].append(skill)
    
    # Projects
    project_rows = conn.execute(
        'SELECT * FROM projects WHERE consultant_id = ? ORDER BY display_order, start_date DESC',
        (consultant_id,)
    ).fetchall()

    projects = []
    for row in project_rows:
        project_data = dict(row)
        skills = conn.execute('''
            SELECT s.skill_name
            FROM skills s
            JOIN project_skills ps ON s.id = ps.skill_id
            WHERE ps.project_id = ?
            ORDER BY s.skill_name
        ''', (row['id'],)).fetchall()
        project_data['skills'] = [s['skill_name'] for s in skills]
        projects.append(project_data)
    
    return {
        'personal_info': personal_info,
        'work_experiences': work_experiences,
        'education': education,
        'certifications': certifications,
        'skills_by_category': skills_by_category,
        'projects': projects,
        't': TRANSLATIONS.get(session.get('language', 'en'), TRANSLATIONS['en']),
        'current_lang': session.get('language', 'en')
    }


if __name__ == '__main__':
    init_database()
    app.run(debug=True, host='0.0.0.0', port=5000)

